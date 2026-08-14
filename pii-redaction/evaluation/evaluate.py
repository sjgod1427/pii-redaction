#!/usr/bin/env python3
"""Score the redactor against the hand-annotated gold sample.

Two questions are being answered, and they need different measurements:

* **Recall / precision on annotated blocks** — the classic span-level scores,
  computed on a stratified sample that was drawn *without* looking at what the
  tool found (see build_sample.py), so it can show misses as well as hits.
* **Leakage on the whole document** — every value the tool did redact somewhere
  must be absent from the final .docx everywhere.  A per-block score can look
  excellent while one stray mention on page 300 undoes the whole exercise.

A detection counts as correct when it overlaps a gold span of the same type
(partial credit, because "Pune - 411 004, Maharashtra" and the same string
without the state are both valid address boundaries).  A stricter exact-boundary
score is reported alongside it.  Token-level accuracy is the fraction of
whitespace tokens in the sample whose redacted/kept status matches the gold.

    python evaluation/evaluate.py "../Red Herring Prospectus.docx" \
        --redacted "output/Red Herring Prospectus - REDACTED.docx"
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from collections import defaultdict
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from pii_redactor import Policy  # noqa: E402
from pii_redactor.docx_io import DocxFile  # noqa: E402
from pii_redactor.pipeline import Redactor, _table_rows  # noqa: E402
from pii_redactor.types import Span  # noqa: E402


def load_jsonl(path: str) -> list[dict]:
    return [json.loads(line) for line in open(path, encoding="utf-8") if line.strip()]


def gold_spans(text: str, entities: list[list[str]]) -> list[Span]:
    """Turn (label, exact substring) annotations into character spans."""
    spans: list[Span] = []
    used: list[tuple[int, int]] = []
    for label, snippet in entities:
        pattern = re.escape(snippet).replace(r"\ ", r"\s+")
        for match in re.finditer(pattern, text):
            if any(match.start() < e and s < match.end() for s, e in used):
                continue
            used.append((match.start(), match.end()))
            spans.append(Span(match.start(), match.end(), label, match.group(0), "gold"))
            break
        else:
            raise SystemExit(f"gold annotation not found in block text: {snippet!r}")
    return spans


def overlap(a: Span, b: Span) -> bool:
    return a.start < b.end and b.start < a.end


def score(predicted: list[Span], gold: list[Span], type_sensitive: bool = True, exact: bool = False):
    """Greedy one-to-one matching between predictions and gold spans."""
    matched_gold: set[int] = set()
    true_positives = 0
    matched_pred: set[int] = set()
    for p_index, prediction in enumerate(predicted):
        for g_index, truth in enumerate(gold):
            if g_index in matched_gold:
                continue
            if type_sensitive and prediction.label != truth.label:
                continue
            hit = (
                (prediction.start, prediction.end) == (truth.start, truth.end)
                if exact
                else overlap(prediction, truth)
            )
            if hit:
                matched_gold.add(g_index)
                matched_pred.add(p_index)
                true_positives += 1
                break
    false_positives = len(predicted) - len(matched_pred)
    false_negatives = len(gold) - len(matched_gold)
    return true_positives, false_positives, false_negatives


def prf(tp: int, fp: int, fn: int) -> tuple[float, float, float]:
    precision = tp / (tp + fp) if tp + fp else 1.0
    recall = tp / (tp + fn) if tp + fn else 1.0
    f1 = 2 * precision * recall / (precision + recall) if precision + recall else 0.0
    return precision, recall, f1


def token_accuracy(text: str, predicted: list[Span], gold: list[Span]) -> tuple[int, int]:
    """Fraction of tokens whose redacted/not-redacted status is right."""
    def covered(spans: list[Span], start: int, end: int) -> bool:
        return any(s.start < end and start < s.end for s in spans)

    correct = total = 0
    for match in re.finditer(r"\S+", text):
        total += 1
        if covered(predicted, match.start(), match.end()) == covered(gold, match.start(), match.end()):
            correct += 1
    return correct, total


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("document")
    parser.add_argument("--redacted", default="output/Red Herring Prospectus - REDACTED.docx")
    parser.add_argument("--gold", default="evaluation/gold_spans.jsonl")
    parser.add_argument("--sample", default="evaluation/sample.jsonl")
    parser.add_argument("--json-out", default="evaluation/metrics.json")
    parser.add_argument("--model", default="en_core_web_lg")
    parser.add_argument("--no-ner", action="store_true")
    args = parser.parse_args()

    sample = {row["block_id"]: row for row in load_jsonl(args.sample)}
    gold_rows = load_jsonl(args.gold)

    # Run the same two-pass pipeline the redactor uses, so the gazetteer sees
    # the whole document exactly as it does in production.
    policy = Policy(spacy_model=args.model, disable_ner=args.no_ner)
    redactor = Redactor(policy)
    docx_file = DocxFile(args.document)
    texts = [paragraph.text for paragraph in docx_file.paragraphs]
    ner_results = list(redactor.ner.pipe(texts))
    for text in texts:
        redactor.gazetteer.observe_vocabulary(text)
    for text, spans in zip(texts, ner_results):
        redactor.gazetteer.learn_block(text, spans)
    for header, cells in _table_rows(docx_file):
        redactor.gazetteer.learn_table_row(header, cells)
    redactor.gazetteer.finalise()
    redactor.surrogates.bind(redactor.gazetteer)

    per_type: dict[str, list[int]] = defaultdict(lambda: [0, 0, 0])
    totals = [0, 0, 0]
    totals_exact = [0, 0, 0]
    totals_untyped = [0, 0, 0]
    tokens_correct = tokens_total = 0
    misses: list[tuple[int, str, str]] = []
    spurious: list[tuple[int, str, str]] = []

    for row in gold_rows:
        block_id = row["block_id"]
        text = sample[block_id]["text"]
        paragraph = docx_file.paragraphs[block_id]
        predicted = redactor.detect(text, paragraph.section, ner_results[block_id])
        truth = gold_spans(text, row["entities"])

        for label in {s.label for s in predicted} | {s.label for s in truth}:
            tp, fp, fn = score(
                [s for s in predicted if s.label == label],
                [s for s in truth if s.label == label],
            )
            bucket = per_type[label]
            bucket[0] += tp
            bucket[1] += fp
            bucket[2] += fn

        for index, (sensitive, exact, bucket) in enumerate(
            ((True, False, totals), (True, True, totals_exact), (False, False, totals_untyped))
        ):
            tp, fp, fn = score(predicted, truth, type_sensitive=sensitive, exact=exact)
            bucket[0] += tp
            bucket[1] += fp
            bucket[2] += fn

        correct, total = token_accuracy(text, predicted, truth)
        tokens_correct += correct
        tokens_total += total

        for truth_span in truth:
            if not any(overlap(p, truth_span) for p in predicted):
                misses.append((block_id, truth_span.label, truth_span.text))
        for prediction in predicted:
            if not any(overlap(prediction, t) for t in truth):
                spurious.append((block_id, prediction.label, prediction.text))

    # ---- document-wide leakage check --------------------------------------
    leakage = check_leakage(args.document, args.redacted)
    image_leakage = check_image_leakage(args.document, args.redacted)

    results = {
        "sample": {
            "blocks": len(gold_rows),
            "gold_entities": sum(len(r["entities"]) for r in gold_rows),
            "characters": sum(len(sample[r["block_id"]]["text"]) for r in gold_rows),
        },
        "overall": _fmt(*totals),
        "overall_exact_boundary": _fmt(*totals_exact),
        "overall_type_insensitive": _fmt(*totals_untyped),
        "token_accuracy": round(tokens_correct / tokens_total, 4) if tokens_total else None,
        "per_type": {
            label: _fmt(*counts) for label, counts in sorted(per_type.items())
        },
        "false_negatives": [{"block": b, "type": t, "text": x[:90]} for b, t, x in misses],
        "false_positives": [{"block": b, "type": t, "text": x[:90]} for b, t, x in spurious],
        "document_leakage": leakage,
        "image_leakage": image_leakage,
    }

    Path(args.json_out).parent.mkdir(parents=True, exist_ok=True)
    Path(args.json_out).write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({k: v for k, v in results.items() if k not in ("false_negatives", "false_positives")}, indent=2))
    print(f"\nfalse negatives ({len(misses)}):")
    for block, label, text in misses:
        print(f"  block {block:5d}  {label:12s} {text[:70]}")
    print(f"\nfalse positives ({len(spurious)}):")
    for block, label, text in spurious:
        print(f"  block {block:5d}  {label:12s} {text[:70]}")
    return 0


def _fmt(tp: int, fp: int, fn: int) -> dict:
    precision, recall, f1 = prf(tp, fp, fn)
    return {
        "tp": tp,
        "fp": fp,
        "fn": fn,
        "precision": round(precision, 4),
        "recall": round(recall, 4),
        "f1": round(f1, 4),
    }


def check_leakage(source: str, redacted: str) -> dict:
    """Every value that was replaced somewhere must be gone everywhere."""
    if not Path(redacted).exists():
        return {"status": "redacted document not found"}

    mapping_path = Path("output/mapping.json")
    if not mapping_path.exists():
        return {"status": "mapping.json not found — run redact.py with --mapping"}

    mapping = json.loads(mapping_path.read_text(encoding="utf-8"))
    text = _plain_text(redacted)
    leaks = []
    checked = 0
    for label, pairs in mapping.items():
        for original in pairs:
            if len(original) < 5:
                continue
            checked += 1
            hits = len(re.findall(re.escape(original).replace(r"\ ", r"\s+"), text))
            if hits:
                leaks.append({"type": label, "value": original[:80], "occurrences": hits})
    return {
        "values_checked": checked,
        "values_still_present": len(leaks),
        "residual_rate": round(len(leaks) / checked, 4) if checked else 0.0,
        "examples": leaks[:15],
    }


def check_image_leakage(source: str, redacted: str) -> dict:
    """Read the pictures in the *output* back and look for PII that survived.

    The text-level check cannot see inside a picture, so a scanned ID card would
    pass it while leaking everything.  This re-runs OCR and barcode decoding over
    the images of the finished document and searches the recovered text for the
    values the tool claims to have replaced — the same standard the text is held
    to, applied to the pixels.
    """
    if not Path(redacted).exists():
        return {"status": "redacted document not found"}
    try:
        import cv2
        import numpy
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        return {"status": "OCR/vision engine unavailable — image leak check skipped"}

    mapping_path = Path("output/mapping.json")
    mapping = json.loads(mapping_path.read_text(encoding="utf-8")) if mapping_path.exists() else {}
    originals = [
        original
        for pairs in mapping.values()
        for original in pairs
        if len(original) >= 5
    ]

    engine, detector = RapidOCR(), cv2.QRCodeDetector()
    images, leaks, recovered = 0, [], []
    with zipfile.ZipFile(redacted) as archive:
        for entry in sorted(archive.namelist()):
            if "/media/" not in entry:
                continue
            images += 1
            blob = archive.read(entry)
            found = ""
            try:
                result, _ = engine(blob)
                found = " ".join(str(item[1]) for item in (result or []) if len(item) > 1)
            except Exception:
                pass
            try:
                array = cv2.imdecode(numpy.frombuffer(blob, numpy.uint8), cv2.IMREAD_COLOR)
                if array is not None:
                    payload = detector.detectAndDecode(array)[0]
                    if payload:
                        found = f"{found} {payload}"
            except Exception:
                pass
            recovered.append({"image": entry, "text": " ".join(found.split())[:120]})
            haystack = _squash(found)
            for original in originals:
                if len(original) >= 5 and _squash(original) in haystack:
                    leaks.append({"image": entry, "value": original[:80]})
    return {
        "images_checked": images,
        "values_still_present": len(leaks),
        "examples": leaks[:15],
        "recovered_text": recovered,
    }


def _squash(text: str) -> str:
    """Casefold and drop non-alphanumerics — OCR loses spaces and punctuation."""
    return re.sub(r"[^a-z0-9]", "", text.casefold())


def _plain_text(path: str) -> str:
    import docx

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    document = docx.Document(path)
    parts = []
    for paragraph in document.element.body.iter(f"{W}p"):
        parts.append("".join(node.text or "" for node in paragraph.iter(f"{W}t")))
    return "\n".join(parts)


if __name__ == "__main__":
    raise SystemExit(main())
