"""Reading text out of a .docx and writing redacted text back in place.

Word stores a paragraph as a sequence of runs, and a single word is routinely
split across several runs (spell-check state, formatting, tracked changes).  So
detection runs on the *reconstructed* paragraph text and replacements are mapped
back onto the individual ``<w:t>`` nodes, which keeps every bit of formatting,
numbering, table structure and styling intact.

Coverage includes the body, tables (any nesting depth), headers, footers,
footnotes/endnotes, text boxes, hyperlink targets and document metadata.
"""

from __future__ import annotations

from dataclasses import dataclass, field

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
V = "{urn:schemas-microsoft-com:vml}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


@dataclass
class _Piece:
    node: object | None  # lxml element for editable <w:t>, None for tab/break
    text: str
    editable: bool


@dataclass
class Paragraph:
    """One editable paragraph, flattened to text with an offset map."""

    element: object
    section: str = ""
    #: which part of the package this paragraph lives in (body, header, footnotes…)
    part: str = "body"
    pieces: list[_Piece] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "".join(piece.text for piece in self.pieces)

    def _offsets(self) -> list[tuple[int, int, _Piece]]:
        out, cursor = [], 0
        for piece in self.pieces:
            out.append((cursor, cursor + len(piece.text), piece))
            cursor += len(piece.text)
        return out

    def apply(self, edits: list[tuple[int, int, str]]) -> int:
        """Apply (start, end, replacement) edits; returns how many landed."""
        applied = 0
        for start, end, replacement in sorted(edits, key=lambda e: e[0], reverse=True):
            if self._apply_one(start, end, replacement):
                applied += 1
        return applied

    def _apply_one(self, start: int, end: int, replacement: str) -> bool:
        touched = [
            (p_start, p_end, piece)
            for p_start, p_end, piece in self._offsets()
            if p_start < end and start < p_end
        ]
        editable = [item for item in touched if item[2].editable]
        if not editable:
            return False

        first_start, first_end, first_piece = editable[0]
        head = first_piece.text[: max(0, start - first_start)]
        tail = first_piece.text[max(0, end - first_start) :] if end <= first_end else ""
        first_piece.text = head + replacement + tail
        self._write(first_piece)

        for piece_start, piece_end, piece in editable[1:]:
            if not piece.editable:
                continue
            keep_tail = piece.text[max(0, end - piece_start) :] if end < piece_end else ""
            piece.text = keep_tail
            self._write(piece)
        return True

    @staticmethod
    def _write(piece: _Piece) -> None:
        if piece.node is None:
            return
        piece.node.text = piece.text
        piece.node.set(XML_SPACE, "preserve")


def _closest_paragraph(node) -> object | None:
    parent = node.getparent()
    while parent is not None:
        if parent.tag == f"{W}p":
            return parent
        parent = parent.getparent()
    return None


def _cell_text(tc) -> str:
    return " ".join("".join(t.text or "" for t in p.iter(f"{W}t")) for p in tc.findall(f"{W}p")).strip()


def _table_headers(root) -> dict:
    """Map paragraph element -> header row of its innermost enclosing table.

    Keyed by the element object itself, not ``id(element)``: lxml creates proxy
    objects on demand and recycles them, so an id captured here can belong to a
    different node by the time the paragraphs are walked.  Holding the element
    as the key keeps its proxy alive and the lookup honest.
    """
    mapping: dict = {}
    for table in root.iter(f"{W}tbl"):
        rows = table.findall(f"{W}tr")
        if not rows:
            continue
        cells = rows[0].findall(f"{W}tc")
        header = " | ".join(_cell_text(tc) for tc in cells)[:200]
        if not header:
            continue
        for paragraph in table.iter(f"{W}p"):
            mapping[paragraph] = header
    return mapping


@dataclass
class ImageRef:
    """One embedded image part, plus every place the document draws it."""

    part: object
    nodes: list = field(default_factory=list)
    contexts: list[str] = field(default_factory=list)
    frozen: str = ""

    @property
    def name(self) -> str:
        return str(getattr(self.part, "partname", "image"))

    @property
    def context(self) -> str:
        """Text surrounding the picture — the signal for signature blocks."""
        if self.frozen:
            return self.frozen
        seen, parts = set(), []
        for chunk in self.contexts:
            cleaned = " ".join(chunk.split())
            if cleaned and cleaned not in seen:
                seen.add(cleaned)
                parts.append(cleaned)
        return "  ".join(parts)[:600]

    def freeze(self) -> "ImageRef":
        """Snapshot the surrounding text *before* it gets redacted.

        A logo's caption is the best clue to which company it belongs to, but by
        the time pictures are processed the caption has already been replaced by
        a surrogate the gazetteer has never heard of.  Capturing it up front is
        what lets a logo inherit the same fake name the prose received.
        """
        self.frozen = self.context
        return self


def _paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(f"{W}t"))


def _picture_context(node) -> str:
    """Text of the paragraph holding a picture, plus its immediate neighbours.

    A signature is usually captioned ("Signature", "Authorised Signatory",
    "For and on behalf of ..."), and that caption is frequently in the adjacent
    paragraph rather than the one holding the image.
    """
    paragraph = _closest_paragraph(node)
    if paragraph is None:
        return ""
    chunks = [_paragraph_text(paragraph)]
    parent = paragraph.getparent()
    if parent is not None:
        siblings = [child for child in parent if child.tag == f"{W}p"]
        try:
            index = siblings.index(paragraph)
        except ValueError:
            index = -1
        if index >= 0:
            for offset in (-2, -1, 1, 2):
                neighbour = index + offset
                if 0 <= neighbour < len(siblings):
                    chunks.append(_paragraph_text(siblings[neighbour]))
    return " ".join(chunk for chunk in chunks if chunk.strip())


def _build_paragraph(element, section: str) -> Paragraph:
    paragraph = Paragraph(element=element, section=section)
    for node in element.iter():
        if _closest_paragraph(node) is not element:
            continue  # belongs to a nested paragraph (e.g. inside a text box)
        if node.tag == f"{W}t":
            paragraph.pieces.append(_Piece(node, node.text or "", True))
        elif node.tag == f"{W}tab":
            paragraph.pieces.append(_Piece(None, "\t", False))
        elif node.tag in (f"{W}br", f"{W}cr"):
            paragraph.pieces.append(_Piece(None, "\n", False))
    return paragraph


class DocxFile:
    """Loads a document and exposes every paragraph of every part."""

    def __init__(self, path: str) -> None:
        self.path = path
        self.document = docx.Document(path)
        self.paragraphs: list[Paragraph] = []
        self._field_nodes: list[object] = []
        self._collect()

    def _collect(self) -> None:
        for root, part in self._roots():
            headers = _table_headers(root)
            for element in root.iter(f"{W}p"):
                paragraph = _build_paragraph(element, headers.get(element, ""))
                if paragraph.pieces:
                    paragraph.part = part
                    self.paragraphs.append(paragraph)
            self._field_nodes.extend(root.iter(f"{W}instrText"))

    def _roots(self):
        yield self.document.element.body, "body"
        for section in self.document.sections:
            for part, name in (
                (section.header, "header"), (section.footer, "footer"),
                (section.even_page_header, "header:even"), (section.even_page_footer, "footer:even"),
                (section.first_page_header, "header:first"), (section.first_page_footer, "footer:first"),
            ):
                try:
                    if part is not None:
                        yield part._element, name
                except Exception:  # pragma: no cover - optional parts
                    continue
        package = self.document.part.package
        for part in package.iter_parts():
            if any(kind in part.partname for kind in ("footnotes", "endnotes", "comments")):
                element = getattr(part, "element", None) or getattr(part, "_element", None)
                if element is not None:
                    yield element, str(part.partname).rsplit("/", 1)[-1].split(".")[0]

    # -- hyperlinks and metadata -------------------------------------------
    def rewrite_hyperlinks(self, rewrite) -> int:
        """Rewrite external hyperlink targets (mailto:/http) via ``rewrite``."""
        changed = 0
        package = self.document.part.package
        for part in package.iter_parts():
            rels = getattr(part, "rels", None)
            if rels is None:
                continue
            for rel in list(rels.values()):
                if rel.reltype != RT.HYPERLINK or not rel.is_external:
                    continue
                new_target = rewrite(rel.target_ref)
                if new_target and new_target != rel.target_ref:
                    rel._target = new_target
                    changed += 1
        return changed

    def field_texts(self) -> list[object]:
        return self._field_nodes

    # -- embedded pictures --------------------------------------------------
    def image_references(self) -> list["ImageRef"]:
        """Every embedded picture, grouped by the image part it points at.

        Walking *parts* rather than the body means headers, footers, footnotes,
        text boxes and any nesting depth are covered without special cases —
        a picture is found wherever Word happened to put it.  Both DrawingML
        (``<a:blip r:embed>``) and legacy VML (``<v:imagedata r:id>``) are
        handled, since documents converted from .doc still use the latter.
        """
        by_part: dict[object, ImageRef] = {}
        package = self.document.part.package
        for part in package.iter_parts():
            element = getattr(part, "element", None)
            if element is None:
                element = getattr(part, "_element", None)
            rels = getattr(part, "rels", None)
            if element is None or rels is None:
                continue
            for node in element.iter():
                rel_id = None
                if node.tag == f"{A}blip":
                    rel_id = node.get(f"{R}embed") or node.get(f"{R}link")
                elif node.tag == f"{V}imagedata":
                    rel_id = node.get(f"{R}id")
                if not rel_id:
                    continue
                try:
                    relationship = rels[rel_id]
                    if relationship.is_external:
                        continue
                    image_part = relationship.target_part
                except (KeyError, ValueError, AttributeError):
                    continue
                if not hasattr(image_part, "blob"):
                    continue
                ref = by_part.get(image_part)
                if ref is None:
                    ref = ImageRef(part=image_part)
                    by_part[image_part] = ref
                ref.nodes.append(node)
                ref.contexts.append(_picture_context(node))
        return list(by_part.values())

    @staticmethod
    def remove_picture(ref: "ImageRef") -> int:
        """Delete the picture itself — the fallback when bytes can't be rewritten.

        Some embeddings (EMF/WMF vector metafiles, SVG) cannot be re-encoded by
        Pillow.  Deleting the drawing removes the content entirely, which is
        always a safe redaction even when replacing it is not possible.
        """
        removed = 0
        for node in ref.nodes:
            container = node
            while container is not None:
                if container.tag in (f"{W}drawing", f"{W}pict", f"{W}object"):
                    break
                container = container.getparent()
            target = container if container is not None else node
            parent = target.getparent()
            if parent is not None:
                parent.remove(target)
                removed += 1
        return removed

    def scrub_metadata(self) -> None:
        props = self.document.core_properties
        props.author = "Redacted"
        props.last_modified_by = "Redacted"
        props.comments = ""
        props.category = ""
        props.keywords = ""
        for attribute in ("title", "subject"):
            if getattr(props, attribute, None):
                setattr(props, attribute, "Redacted document")

    def save(self, path: str) -> None:
        self.document.save(path)
