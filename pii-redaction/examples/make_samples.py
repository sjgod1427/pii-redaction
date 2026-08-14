#!/usr/bin/env python3
"""Build the sample documents offered in the web UI's library.

These are deliberately long and verbose — a few hundred paragraphs across a
dozen sections, with tables, schedules and embedded images. Short fixtures make
a redaction tool look better than it is: the document-level gazetteer only earns
its keep when a name introduced in section 2 has to be recognised again in an
all-caps table cell in section 11, and precision problems only show up once
there is enough surrounding prose to trip over.

Every value in every document here is invented.

    python examples/make_samples.py
"""

from __future__ import annotations

import io
import math
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE.parent))


# --- image builders ---------------------------------------------------------
def _id_card(name, parent, dob, number, address):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (560, 340), "#eef3fb")
    draw = ImageDraw.Draw(image)
    draw.rectangle([(0, 0), (559, 44)], fill="#2f4f8f")
    draw.text((16, 18), "NATIONAL IDENTITY AUTHORITY", fill="white")
    for index, line in enumerate(
        [f"Name: {name}", f"Father's Name: {parent}", f"Date of Birth: {dob}",
         f"Account Number: {number}", f"Address: {address}"]
    ):
        draw.text((176, 70 + index * 34), line, fill="#101820")
    draw.rectangle([(20, 66), (156, 250)], fill="#c8d4e6")
    draw.ellipse([(56, 96), (120, 168)], fill="#8b7355")
    draw.ellipse([(40, 176), (136, 268)], fill="#4a5b78")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _signature(seed=3):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (420, 130), "white")
    draw = ImageDraw.Draw(image)
    points = [
        (30 + step * 3.6, 84 - math.sin(step / (6.0 + seed)) * 26 - math.sin(step / 2.3) * 7)
        for step in range(100)
    ]
    draw.line(points, fill="#1b2c5a", width=3, joint="curve")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _logo(text, colour):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (300, 110), "white")
    draw = ImageDraw.Draw(image)
    draw.ellipse([(12, 24), (74, 86)], fill=colour)
    draw.text((92, 50), text, fill="#12233a")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _qr(payload):
    import qrcode

    buffer = io.BytesIO()
    qrcode.make(payload).save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


# --- document helpers -------------------------------------------------------
class Doc:
    """Thin wrapper that keeps the section/clause numbering tidy."""

    def __init__(self, title, subtitle):
        self.document = docx.Document()
        self.section_no = 0
        self.document.add_heading(title, level=0)
        para = self.document.add_paragraph(subtitle)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def section(self, title):
        self.section_no += 1
        self.clause_no = 0
        self.document.add_heading(f"{self.section_no}. {title}", level=1)
        return self.section_no

    def clause(self, text):
        self.clause_no += 1
        self.document.add_paragraph(f"{self.section_no}.{self.clause_no}  {text}")

    def para(self, text):
        self.document.add_paragraph(text)

    def bullets(self, items):
        for item in items:
            self.document.add_paragraph(item, style="List Bullet")

    def table(self, header, rows):
        table = self.document.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        for index, label in enumerate(header):
            table.rows[0].cells[index].text = label
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
        self.document.add_paragraph("")

    def picture(self, stream, inches):
        self.document.add_picture(stream, width=Inches(inches))

    def save(self, path):
        self.document.save(path)
        blocks = len(self.document.element.body.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"))
        print(f"wrote {path}  ({blocks} blocks)")


# --- 1. employment contract -------------------------------------------------
def build_offer_letter(path: Path) -> None:
    doc = Doc("Employment Agreement and Onboarding File",
              "Kestrel Dynamics Private Limited · Reference OFR/2026/00184 · Confidential")
    doc.picture(_logo("Kestrel Dynamics", "#2f6bd8"), 1.9)

    doc.section("Parties and recitals")
    doc.clause(
        "This Employment Agreement is made at Navi Mumbai on the 14th day of January 2026 between "
        "Kestrel Dynamics Private Limited, a company incorporated under the Companies Act, 2013, "
        "having its registered office at 4th Floor, Orion Tower, Sector 14, Vashi, Navi Mumbai – 400 703, "
        "Maharashtra, India (hereinafter the \"Company\"), of the first part;")
    doc.clause(
        "And Anwesha Bhattacharya, daughter of Sujoy Bhattacharya, residing at Flat 702, Sunbeam Residency, "
        "Kalyani Nagar, Pune – 411 006, Maharashtra, India (hereinafter the \"Employee\"), of the second part.")
    doc.clause(
        "The Company is engaged in the design and manufacture of precision instrumentation and has "
        "manufacturing facilities at Vashi, Ranjangaon and Hosur. The Employee has represented that she "
        "possesses the qualifications and experience necessary for the role offered.")
    doc.clause(
        "The Employee was interviewed by Devendra Raghunath Kulkarni, Vice President of Engineering, and by "
        "Farida Anwar Sheikh, Head of Talent, on 22 December 2025 and 29 December 2025 respectively.")
    doc.clause(
        "The parties have agreed to record the terms of employment in writing. This Agreement supersedes the "
        "offer summary emailed to anwesha.bhattacharya@gmail.com on 5 January 2026.")

    doc.section("Personal particulars of the Employee")
    doc.para("The following particulars have been verified against original documents produced by the Employee.")
    doc.table(
        ["Field", "Detail"],
        [
            ("Full name", "Anwesha Bhattacharya"),
            ("Father's name", "Sujoy Bhattacharya"),
            ("Date of birth", "17 February 1994"),
            ("Personal email", "anwesha.bhattacharya@gmail.com"),
            ("Work email", "anwesha.b@kestreldynamics.com"),
            ("Mobile", "+91 98204 11762"),
            ("Alternate contact", "+91 20 6644 8891"),
            ("Residential address", "Flat 702, Sunbeam Residency, Kalyani Nagar, Pune – 411 006, Maharashtra, India"),
            ("Permanent address", "12 Copperfield Lane, Salt Lake Sector II, Kolkata – 700 091, West Bengal, India"),
            ("PAN", "AVKPB7391L"),
            ("Aadhaar", "6621 4470 8853"),
            ("Passport number", "M4471209"),
            ("Bank account number", "50100248871904"),
            ("IFSC", "HDFC0000521"),
            ("Provident fund UAN", "100987654321"),
            ("Employee ID", "KD-2291"),
            ("Emergency contact", "Sujoy Bhattacharya, +91 98304 22187"),
        ],
    )
    doc.para("Proof of identity submitted with the application:")
    doc.picture(_id_card("ANWESHA BHATTACHARYA", "Sujoy Bhattacharya", "17/02/1994",
                         "AVKPB7391L", "702 Sunbeam Residency, Pune - 411 006"), 3.4)

    doc.section("Appointment, designation and reporting")
    for text in [
        "The Employee is appointed as Senior Reliability Engineer, Grade M3, in the Engineering Quality function.",
        "The Employee shall report to Devendra Raghunath Kulkarni, Vice President of Engineering, or to such "
        "other person as the Company may designate from time to time.",
        "The Employee's dotted-line reporting for quality matters shall be to Meera Kandaswamy Pillai, "
        "General Manager, Reliability Assurance.",
        "The place of posting shall be the Company's Navi Mumbai office, with an expectation of travel to the "
        "Ranjangaon and Hosur plants of up to eight days per month.",
        "The Company may transfer the Employee to any of its offices, plants, subsidiaries or associate "
        "companies, including Kestrel Instruments LLP and Orion Testing Services Private Limited.",
        "The appointment is effective from 2 February 2026 and is subject to the satisfactory completion of "
        "the verification described in Section 9.",
    ]:
        doc.clause(text)

    doc.section("Remuneration and benefits")
    doc.para("The annual fixed compensation is set out below. All amounts are in Indian Rupees.")
    doc.table(
        ["Component", "Per annum", "Per month", "Notes"],
        [
            ("Basic salary", "9,60,000", "80,000", "40% of fixed pay"),
            ("House rent allowance", "4,80,000", "40,000", "50% of basic"),
            ("Special allowance", "6,72,000", "56,000", "Balancing component"),
            ("Leave travel allowance", "60,000", "5,000", "Claim against proof"),
            ("Provident fund (employer)", "1,15,200", "9,600", "12% of basic"),
            ("Gratuity provision", "46,150", "3,846", "As per Payment of Gratuity Act"),
            ("Medical insurance premium", "38,000", "3,167", "Family floater, 10,00,000 cover"),
            ("Total fixed cost to company", "23,71,350", "1,97,613", ""),
        ],
    )
    for text in [
        "A performance bonus of up to 15% of fixed pay is payable annually, at the sole discretion of the "
        "Company, subject to the Employee being on the rolls on the date of disbursement.",
        "Salary shall be credited to bank account 50100248871904 held with HDFC Bank Limited, Kalyani Nagar "
        "branch, IFSC HDFC0000521, on or before the last working day of each calendar month.",
        "Payroll queries should be addressed to payroll@kestreldynamics.com or to Sunita Ravikant Joshi, "
        "Payroll Manager, on +91 22 6644 1180.",
        "Reimbursements are administered through the internal portal at "
        "https://portal.kestreldynamics.com/expenses and must be submitted within 45 days.",
    ]:
        doc.clause(text)

    doc.section("Working hours, leave and attendance")
    doc.para(
        "The standard working week is 40 hours, Monday to Friday, between 09:30 and 18:30 with a one-hour "
        "break. Attendance is recorded by access card KD-2291 at all Company premises.")
    doc.table(
        ["Leave type", "Entitlement per year", "Carry forward", "Encashable"],
        [
            ("Privilege leave", "21 days", "Up to 45 days", "Yes, on separation"),
            ("Casual leave", "7 days", "Nil", "No"),
            ("Sick leave", "10 days", "Up to 30 days", "No"),
            ("Maternity leave", "26 weeks", "Not applicable", "No"),
            ("Paternity leave", "10 days", "Not applicable", "No"),
            ("Bereavement leave", "5 days", "Nil", "No"),
            ("Public holidays", "11 days", "Nil", "No"),
        ],
    )
    for text in [
        "Leave must be applied for through the human resources portal and approved by the reporting manager "
        "at least five working days in advance, except in the case of sick leave.",
        "Absence exceeding three consecutive working days without approval shall be treated as absence "
        "without leave and may result in disciplinary action under Section 11.",
        "The Company observes the holiday calendar published by Farida Anwar Sheikh each December.",
    ]:
        doc.clause(text)

    doc.section("Confidentiality and intellectual property")
    for text in [
        "The Employee shall not, during the term of employment or at any time thereafter, disclose to any "
        "person any confidential information belonging to the Company, its customers or its suppliers.",
        "Confidential information includes designs, drawings, test data, pricing, customer lists, supplier "
        "terms, unpublished financial information and the personal data of other employees.",
        "All intellectual property created by the Employee in the course of employment shall vest absolutely "
        "in the Company, and the Employee shall execute such documents as may be required to give effect to "
        "this clause.",
        "The Employee shall not retain copies of Company documents on personal devices or personal cloud "
        "storage. Company data on the Employee's laptop, asset tag KD-LAP-4471, is subject to remote wipe.",
        "The obligations in this Section survive the termination of employment for a period of five years, "
        "and indefinitely in respect of trade secrets.",
    ]:
        doc.clause(text)

    doc.section("Non-solicitation and conflict of interest")
    for text in [
        "For a period of twelve months following separation, the Employee shall not solicit any employee of "
        "the Company to leave its employment.",
        "The Employee shall not, without prior written consent, engage in any business which competes with "
        "the Company, including any engagement with Meridian Instruments Limited or Silverbrook Metrology "
        "Private Limited.",
        "The Employee has disclosed a shareholding of 1,200 equity shares in Northgate Analytics Limited, "
        "which the Company has reviewed and does not consider a conflict.",
        "Any change in the disclosures made under this Section must be notified to compliance@kestreldynamics.com "
        "within seven days.",
    ]:
        doc.clause(text)

    doc.section("Data protection")
    doc.para(
        "The Company processes the Employee's personal data for payroll, statutory compliance, benefits "
        "administration and workplace safety. The lawful basis for processing is the performance of this "
        "contract and compliance with legal obligations.")
    doc.table(
        ["Category of data", "Purpose", "Retention", "Recipients"],
        [
            ("Identity and contact", "Payroll, communication", "7 years post separation", "Payroll processor"),
            ("Bank and tax", "Salary disbursement", "8 years", "HDFC Bank Limited, tax authorities"),
            ("Health and insurance", "Medical cover", "3 years", "Northaven Assurance Company Limited"),
            ("Attendance and access", "Site security", "2 years", "Facilities team"),
            ("Performance records", "Appraisal", "5 years", "Reporting manager, HR"),
            ("Background verification", "Pre-employment screening", "3 years", "Sentinel Screening LLP"),
        ],
    )
    doc.clause(
        "The Employee may exercise rights of access and correction by writing to the Data Protection Officer, "
        "Rajarshi Nandkumar Pathak, at dpo@kestreldynamics.com.")

    doc.section("Background verification")
    for text in [
        "This appointment is conditional upon satisfactory verification by Sentinel Screening LLP, "
        "reachable at verify@sentinelscreening.co.in or +91 22 6644 1200.",
        "Verification covers identity, address, education, prior employment and criminal record checks.",
        "The Employee's prior employment with Harborview Controls Private Limited from June 2019 to "
        "December 2025 will be verified with its human resources department.",
        "Adverse findings may result in the withdrawal of this offer or termination without notice.",
    ]:
        doc.clause(text)
    doc.para("Scan to confirm acceptance of this offer:")
    doc.picture(_qr("https://kestreldynamics.com/offer/accept/OFR-2026-00184"), 1.1)

    doc.section("Health, safety and insurance")
    doc.para(
        "The Employee is covered under the Company's group medical policy with Northaven Assurance Company "
        "Limited, policy number NAV-GRP-8891204, administered by Fairwind Benefits Services LLP.")
    doc.table(
        ["Cover", "Sum insured", "Dependants covered", "Claims contact"],
        [
            ("Group medical", "10,00,000", "Spouse, 2 children, parents", "claims@northavenassure.com"),
            ("Group personal accident", "24,00,000", "Employee only", "+91 22 6712 4400"),
            ("Group term life", "36,00,000", "Employee only", "Nominee as declared"),
            ("Travel insurance", "USD 250,000", "Employee only", "For overseas travel"),
        ],
    )
    doc.clause(
        "The Employee has nominated Sujoy Bhattacharya, father, as the beneficiary under the group term life "
        "policy, and has provided his Aadhaar 7419 2280 6634 for that purpose.")

    doc.section("Discipline and grievance")
    for text in [
        "The Company maintains a code of conduct published on the intranet. Breach may lead to warning, "
        "suspension, withholding of increment or termination.",
        "Grievances should be raised first with the reporting manager and, if unresolved, with the Head of "
        "Talent, Farida Anwar Sheikh, at farida.sheikh@kestreldynamics.com.",
        "Complaints of sexual harassment are dealt with by the Internal Committee chaired by Meera Kandaswamy "
        "Pillai, in accordance with the Sexual Harassment of Women at Workplace Act, 2013.",
        "Whistleblower reports may be made anonymously at https://kestreldynamics.com/ethics or to the "
        "external ombudsman, Kirtane Advisory LLP.",
    ]:
        doc.clause(text)

    doc.section("Termination")
    for text in [
        "Either party may terminate this Agreement by giving 90 days' written notice, or salary in lieu.",
        "During probation of six months, the notice period shall be 30 days.",
        "The Company may terminate without notice for misconduct, including falsification of records, breach "
        "of confidentiality or conviction of a criminal offence.",
        "On separation the Employee shall return all Company property, including laptop KD-LAP-4471, access "
        "card KD-2291 and any documents in physical or electronic form.",
        "Full and final settlement shall be completed within 45 days of the last working day and credited to "
        "the account recorded in Section 4.",
    ]:
        doc.clause(text)

    doc.section("Governing law and execution")
    doc.clause("This Agreement is governed by the laws of India, subject to the jurisdiction of the courts at Mumbai.")
    doc.clause(
        "Disputes shall first be referred to mediation before Anand Vishwanath Rao, Advocate, of "
        "Rao & Pandit Associates, 18 Fairwind Road, Bandra East, Mumbai – 400 051, Maharashtra, India.")
    doc.para("For and on behalf of Kestrel Dynamics Private Limited")
    doc.para("Authorised Signatory")
    doc.picture(_signature(2), 2.2)
    doc.para("Name: Devendra Raghunath Kulkarni · Designation: Vice President, Engineering")
    doc.para("")
    doc.para("Accepted by the Employee")
    doc.picture(_signature(7), 2.2)
    doc.para("Name: Anwesha Bhattacharya · Date: 14 January 2026")

    doc.section("Annexure A — asset issue and reference numbers")
    doc.para(
        "The following references are internal identifiers. They identify equipment and transactions rather "
        "than people, and are recorded here for completeness.")
    doc.table(
        ["Reference", "Description", "Raised on"],
        [
            ("PO 4500219873", "Laptop procurement", "8 January 2026"),
            ("Invoice 993214", "Laptop, monitor, dock", "9 January 2026"),
            ("Ticket #4482910", "IT asset provisioning", "10 January 2026"),
            ("Order 100002345", "Ergonomic chair", "10 January 2026"),
            ("Asset KD-LAP-4471", "ThinkPad, 32 GB", "12 January 2026"),
            ("Policy revision 3.2", "Travel policy applicable", "4 March 2021"),
            ("Requisition REQ-2026-0091", "Headcount approval", "2 December 2025"),
        ],
    )
    doc.para(
        "System access will be provisioned from the corporate range and first login is expected from "
        "203.0.113.88. Access to the onboarding portal at "
        "https://portal.kestreldynamics.com/onboarding/anwesha-bhattacharya expires after 14 days.")
    doc.save(path)


# --- 2. insurance claim file ------------------------------------------------
def build_claim_form(path: Path) -> None:
    doc = Doc("Health Insurance Claim File",
              "Northaven Assurance Company Limited · Claim CLM/2026/778341 · Confidential")
    doc.picture(_logo("Northaven Assure", "#1f9e86"), 1.9)

    doc.section("Claim summary")
    doc.clause(
        "This file records the intimation, assessment and settlement of claim CLM/2026/778341 submitted to "
        "Northaven Assurance Company Limited, Claims Processing Centre, 22 Fairwind Road, Bandra East, "
        "Mumbai – 400 051, Maharashtra, India.")
    doc.clause(
        "The claim was intimated on 8 January 2026 at 11:42 hours by the policyholder, Ramaswamy Venkatesh "
        "Iyer, by telephone on +91 98867 40219, and confirmed by email from rv.iyer@yahoo.co.in.")
    doc.clause(
        "The claim relates to a planned surgical admission at Meridian Speciality Hospital, Malleshwaram, "
        "Bengaluru, under the care of Dr. Shalini Prabhakar Deshpande.")
    doc.clause("The claim was registered from IP address 198.51.100.42 through the customer portal.")

    doc.section("Policyholder particulars")
    doc.table(
        ["Field", "Detail"],
        [
            ("Policyholder name", "Ramaswamy Venkatesh Iyer"),
            ("Father's name", "Venkatesh Subramanian Iyer"),
            ("Date of birth", "03/09/1968"),
            ("Age at admission", "57 years"),
            ("Policy number", "NAV-HLT-4471209"),
            ("Policy inception", "12 April 2016"),
            ("Aadhaar", "7412 8890 3365"),
            ("PAN", "AXQPI2274K"),
            ("Mobile", "+91 98867 40219"),
            ("Email", "rv.iyer@yahoo.co.in"),
            ("Residential address", "18 Copperfield Lane, Malleshwaram, Bengaluru – 560 003, Karnataka, India"),
            ("Employer", "Silverbrook Metrology Private Limited"),
            ("Bank account for settlement", "0091234567890"),
            ("IFSC", "ICIC0000091"),
            ("Card used for deposit", "4111 1111 1111 1111"),
        ],
    )
    doc.para("Identity proof submitted with the claim:")
    doc.picture(_id_card("RAMASWAMY V IYER", "Venkatesh Iyer", "03/09/1968", "AXQPI2274K",
                         "18 Copperfield Lane, Bengaluru - 560 003"), 3.4)

    doc.section("Insured members under the policy")
    doc.table(
        ["Name", "Relationship", "Date of birth", "Aadhaar", "Sum insured"],
        [
            ("Ramaswamy Venkatesh Iyer", "Self", "03/09/1968", "7412 8890 3365", "10,00,000"),
            ("Lakshmi Ramaswamy Iyer", "Spouse", "22/11/1971", "8830 4471 9926", "10,00,000"),
            ("Aditya Ramaswamy Iyer", "Son", "14/06/2001", "9021 7734 5518", "5,00,000"),
            ("Sharada Venkatesh Iyer", "Mother", "09/02/1944", "6612 8890 3341", "3,00,000"),
        ],
    )
    doc.clause(
        "The nominee on record is Lakshmi Ramaswamy Iyer, spouse, reachable at lakshmi.iyer@gmail.com and "
        "on +91 99456 21180.")

    doc.section("Hospital and treating team")
    doc.table(
        ["Field", "Detail"],
        [
            ("Hospital", "Meridian Speciality Hospital"),
            ("Address", "44 Sampige Road, Malleshwaram, Bengaluru – 560 003, Karnataka, India"),
            ("Network status", "Cashless network provider"),
            ("Hospital registration", "KA/BLR/HOSP/2011/4471"),
            ("Treating physician", "Dr. Shalini Prabhakar Deshpande"),
            ("Registration number", "KMC-88214"),
            ("Anaesthetist", "Dr. Girish Mohan Rane"),
            ("Surgeon", "Dr. Prakash Anantharaman Setty"),
            ("Case manager", "Kavitha Nair"),
            ("Hospital contact", "+91 80 4123 7788"),
            ("Billing contact", "billing@meridianspeciality.in"),
        ],
    )
    doc.clause(
        "Pre-authorisation was discussed with Kavitha Nair on 9 January 2026 and approved for an initial "
        "amount of 2,40,000 pending final billing.")

    doc.section("Clinical narrative")
    for text in [
        "The insured presented with a six-month history of progressive right knee pain, worse on stair "
        "descent, with morning stiffness lasting approximately twenty minutes.",
        "Conservative management including physiotherapy under Ms. Divya Ranganathan and intra-articular "
        "injection administered on 3 October 2025 provided only transient relief.",
        "Radiographs dated 18 December 2025 demonstrated tricompartmental osteoarthritis with medial joint "
        "space narrowing and subchondral sclerosis.",
        "The insured was counselled regarding total knee replacement and elected to proceed. Consent was "
        "obtained in the presence of Lakshmi Ramaswamy Iyer.",
        "Pre-operative assessment recorded blood pressure of 138/86, random blood sugar of 142 mg/dL and no "
        "history of ischaemic heart disease.",
        "The insured is a known hypertensive on Amlodipine 5 mg once daily, prescribed by "
        "Dr. Shalini Prabhakar Deshpande since 2019.",
        "Surgery was performed on 14 January 2026. The post-operative period was uneventful and the insured "
        "was mobilised on day one with a walker.",
        "The insured was discharged on 18 January 2026 with advice for physiotherapy three times weekly for "
        "six weeks and review on 12 February 2026.",
    ]:
        doc.clause(text)

    doc.section("Itemised billing")
    doc.table(
        ["Line", "Description", "Quantity", "Rate", "Amount"],
        [
            ("1", "Room rent — single private, 5 nights", "5", "9,000", "45,000"),
            ("2", "Surgeon's professional fee", "1", "1,10,000", "1,10,000"),
            ("3", "Anaesthetist's fee", "1", "38,000", "38,000"),
            ("4", "Operation theatre charges", "1", "62,000", "62,000"),
            ("5", "Knee prosthesis, cemented", "1", "1,48,000", "1,48,000"),
            ("6", "Pharmacy and consumables", "—", "—", "54,320"),
            ("7", "Investigations — pre-operative", "—", "—", "18,900"),
            ("8", "Physiotherapy, inpatient", "5", "1,200", "6,000"),
            ("9", "Nursing charges", "5", "2,400", "12,000"),
            ("10", "Discharge medication", "—", "—", "4,180"),
            ("", "Gross total", "", "", "4,98,400"),
            ("", "Non-payable items", "", "", "22,150"),
            ("", "Net payable", "", "", "4,76,250"),
        ],
    )
    doc.clause(
        "Hospital bill number 20260114-8871 and pharmacy invoice 55219 are enclosed. Order 100002345 relates "
        "to equipment rental arranged privately and does not form part of this claim.")

    doc.section("Assessment and deductions")
    doc.table(
        ["Head", "Claimed", "Admitted", "Deducted", "Reason"],
        [
            ("Room rent", "45,000", "45,000", "Nil", "Within eligibility"),
            ("Surgeon", "1,10,000", "1,10,000", "Nil", "Within schedule"),
            ("Prosthesis", "1,48,000", "1,48,000", "Nil", "Covered implant"),
            ("Consumables", "54,320", "40,170", "14,150", "Non-medical items"),
            ("Investigations", "18,900", "18,900", "Nil", "Related to admission"),
            ("Attendant charges", "8,000", "Nil", "8,000", "Excluded under policy"),
        ],
    )
    for text in [
        "The claim was assessed by Priyanka Deshmukh Salunkhe, Senior Claims Officer, and reviewed by "
        "Dr. Arvind Kumar Bhatnagar, Panel Medical Officer.",
        "Deductions were communicated to the policyholder by email on 22 January 2026 and acknowledged.",
        "No pre-existing disease exclusion was applied, the policy having completed its waiting period in "
        "April 2020.",
    ]:
        doc.clause(text)

    doc.section("Correspondence log")
    doc.table(
        ["Date", "From", "To", "Channel", "Summary"],
        [
            ("08/01/2026", "Ramaswamy Venkatesh Iyer", "Claims desk", "Telephone", "Intimation of planned admission"),
            ("09/01/2026", "Kavitha Nair", "Priyanka Deshmukh Salunkhe", "Email", "Pre-authorisation request"),
            ("09/01/2026", "Priyanka Deshmukh Salunkhe", "Meridian Speciality Hospital", "Portal", "Initial approval 2,40,000"),
            ("15/01/2026", "Meridian Speciality Hospital", "Claims desk", "Portal", "Enhancement request"),
            ("16/01/2026", "Dr. Arvind Kumar Bhatnagar", "Claims desk", "Internal note", "Enhancement supported"),
            ("18/01/2026", "Meridian Speciality Hospital", "Claims desk", "Courier", "Final bill and discharge summary"),
            ("22/01/2026", "Priyanka Deshmukh Salunkhe", "rv.iyer@yahoo.co.in", "Email", "Deduction advice"),
            ("24/01/2026", "Lakshmi Ramaswamy Iyer", "Claims desk", "Telephone", "Query on attendant charges"),
            ("27/01/2026", "Claims desk", "Meridian Speciality Hospital", "NEFT", "Settlement remitted"),
        ],
    )
    doc.para(
        "Supporting documents were uploaded to "
        "https://claims.northavenassure.com/upload/rv-iyer-778341 and are retained for eight years.")

    doc.section("Policy terms relied upon")
    doc.bullets([
        "Clause 4.2 — room rent eligibility capped at 2% of sum insured per day.",
        "Clause 5.1 — implants and prostheses covered when medically necessary.",
        "Clause 6.7 — non-medical consumables listed in Annexure II are not payable.",
        "Clause 7.3 — attendant and companion charges are excluded.",
        "Clause 9.1 — cashless settlement available at network providers.",
        "Clause 11.4 — claims must be intimated within 48 hours of planned admission.",
        "Clause 12.2 — the insurer may seek an independent medical opinion.",
        "Clause 14.6 — settlement within 30 days of receipt of complete documents.",
    ])

    doc.section("Settlement and declaration")
    doc.clause(
        "The net admitted amount of 4,76,250 was remitted on 27 January 2026 to Meridian Speciality Hospital "
        "under cashless settlement, UTR NAV20260127884412.")
    doc.clause(
        "The policyholder's residual liability of 22,150 was settled directly at discharge by card ending "
        "4111 1111 1111 1111.")
    doc.para(
        "I declare that the information given above is true to the best of my knowledge and that I have not "
        "concealed any material fact.")
    doc.para("Signature of claimant")
    doc.picture(_signature(5), 2.2)
    doc.para("Name: Ramaswamy Venkatesh Iyer · Date: 28 January 2026 · Place: Bengaluru")
    doc.para("Scan for claim status:")
    doc.picture(_qr("https://claims.northavenassure.com/status/CLM-2026-778341"), 1.1)
    doc.save(path)


# --- 3. support case file ---------------------------------------------------
def build_ticket_log(path: Path) -> None:
    doc = Doc("Customer Support Case File",
              "Sunrise Textiles Private Limited · Escalation ESC/2026/0447 · Internal")
    doc.picture(_logo("Sunrise Textiles", "#d63b2f"), 1.9)

    doc.section("Escalation summary")
    doc.clause(
        "This file consolidates twelve support tickets raised between 2 January and 30 January 2026 by "
        "customers of Sunrise Textiles Private Limited, 27 Industrial Estate, Panchvati, Pashan, "
        "Pune – 411 008, Maharashtra, India.")
    doc.clause(
        "The escalation was opened by Priya Nair, Account Manager, following a cluster of payment failures "
        "reported after the billing platform migration of 28 December 2025.")
    doc.clause(
        "The file is circulated to the Customer Operations Committee and to Vertex Logistics LLP as the "
        "fulfilment partner.")

    doc.section("Tickets raised")
    doc.table(
        ["Ticket", "Date", "Customer", "Contact", "Category", "Status"],
        [
            ("TKT-100294", "02/01/2026", "Rashi Patil", "rashhi.patil@gmail.com", "Payment declined", "Closed"),
            ("TKT-100295", "04/01/2026", "Rohan Dey", "rohan.dey@gmail.com", "Order not received", "Closed"),
            ("TKT-100296", "06/01/2026", "Meenakshi Sundaram", "m.sundaram@outlook.com", "Wrong item", "Closed"),
            ("TKT-100297", "09/01/2026", "Imtiaz Rehman Qureshi", "imtiaz.q@rediffmail.com", "Refund pending", "Escalated"),
            ("TKT-100298", "11/01/2026", "Sneha Kulkarni", "sneha.kulkarni@yahoo.co.in", "Duplicate charge", "Closed"),
            ("TKT-100299", "14/01/2026", "Bhaskar Reddy Konda", "bhaskar.konda@gmail.com", "Address change", "Closed"),
            ("TKT-100300", "17/01/2026", "Anita D'Souza", "anita.dsouza@hotmail.com", "Damaged goods", "Escalated"),
            ("TKT-100301", "20/01/2026", "Tarun Vasudev Shenoy", "tarun.shenoy@gmail.com", "Invoice query", "Closed"),
            ("TKT-100302", "22/01/2026", "Fatima Bano Ansari", "fatima.ansari@gmail.com", "Payment declined", "Escalated"),
            ("TKT-100303", "25/01/2026", "Devika Menon", "devika.menon@protonmail.com", "Account locked", "Closed"),
            ("TKT-100304", "28/01/2026", "Harpreet Singh Bedi", "hs.bedi@gmail.com", "Delivery delay", "Open"),
            ("TKT-100305", "30/01/2026", "Nandini Iyer Raghavan", "nandini.raghavan@gmail.com", "Refund pending", "Open"),
        ],
    )

    doc.section("Ticket TKT-100294 — payment declined")
    for text in [
        "Rashi Patil of Sunrise Textiles Private Limited reported that her card ending 4111 1111 1111 1111 "
        "was declined three times on 2 January 2026 while settling invoice 993214.",
        "The customer's registered date of birth is 14 March 1988 and the billing address on file is "
        "12 Buena Monte, Panchvati, Pashan, Pune – 411 008, Maharashtra, India.",
        "The customer was contacted on +91 98765 43210 by Sagar Bhatt of the payments team.",
        "Investigation showed the acquirer had rejected the transaction owing to an address mismatch "
        "following the platform migration.",
        "The billing address was corrected and the payment succeeded on 3 January 2026 against order 100002345.",
        "The ticket was closed on 4 January 2026 with a service credit of 500 applied.",
    ]:
        doc.clause(text)

    doc.section("Ticket TKT-100295 — order not received")
    for text in [
        "Rohan Dey called from 415-555-0132 on 4 January 2026 regarding order 100002345, dispatched on "
        "29 December 2025 through Vertex Logistics LLP.",
        "The customer's SSN on file is 123-45-6789, recorded during the United States pilot programme, and "
        "his address is 940 Larch Street, Springfield, IL 62704, USA.",
        "The last login recorded against the account was from 203.0.113.47 on 2 January 2026.",
        "Tracking showed the consignment had been misrouted to the Ahmednagar hub rather than the "
        "international gateway.",
        "A replacement was dispatched on 7 January 2026 and delivered on 12 January 2026.",
        "The customer confirmed receipt by email and the ticket was closed on 13 January 2026.",
    ]:
        doc.clause(text)

    doc.section("Ticket TKT-100297 — refund pending, escalated")
    for text in [
        "Imtiaz Rehman Qureshi raised a refund request on 9 January 2026 for 18,400 against invoice 993871.",
        "The customer is registered at 4 Nightingale Court, Frazer Town, Bengaluru – 560 005, Karnataka, "
        "India and is reachable on +91 99012 44780.",
        "The refund was initiated on 10 January 2026 but failed owing to a stale beneficiary record with "
        "IFSC ICIC0000091.",
        "Escalated by Priya Nair to Vertex Logistics LLP finance on 15 January 2026.",
        "Revised bank details were collected and verified against the customer's PAN ABCDE1234F.",
        "The refund was reprocessed on 21 January 2026, UTR SUN20260121447901, and remains under confirmation.",
    ]:
        doc.clause(text)

    doc.section("Ticket TKT-100300 — damaged goods, escalated")
    for text in [
        "Anita D'Souza reported damage to a consignment received on 16 January 2026 at "
        "Flat 9, Sea Breeze Apartments, Bandra East, Mumbai – 400 051, Maharashtra, India.",
        "Photographs were submitted through the portal and reviewed by Quality Assurance under "
        "Ganesh Prasad Kulkarni.",
        "The damage was attributed to inadequate palletisation at the Chakan despatch point.",
        "A corrective action request was raised with the warehouse supervisor, Mahesh Tukaram Jadhav.",
        "A full replacement was authorised on 19 January 2026 and the customer was offered a 10% credit.",
        "The ticket remains open pending confirmation of the corrective action.",
    ]:
        doc.clause(text)

    doc.section("Root cause analysis")
    for text in [
        "The billing platform migration of 28 December 2025 did not carry forward normalised address records, "
        "causing address verification failures at the acquirer.",
        "Approximately 340 customers were affected, of whom 12 raised tickets and 4 escalated.",
        "The defect was traced to a field-mapping error in the migration script authored by "
        "Sarthak Malvadkar and reviewed by Anand Soni.",
        "A hotfix was deployed on 12 January 2026 and back-population of affected records completed on "
        "19 January 2026.",
        "No payment card data was exposed at any point; the acquirer's tokenised vault was unaffected.",
    ]:
        doc.clause(text)

    doc.section("Corrective and preventive actions")
    doc.table(
        ["Action", "Owner", "Due", "Status"],
        [
            ("Add address normalisation regression tests", "Sarthak Malvadkar", "31/01/2026", "Complete"),
            ("Migration dry-run on staging mandated", "Anand Soni", "15/02/2026", "In progress"),
            ("Acquirer failure alerting", "Sagar Bhatt", "20/02/2026", "In progress"),
            ("Customer communication template", "Priya Nair", "07/02/2026", "Complete"),
            ("Palletisation audit at Chakan", "Mahesh Tukaram Jadhav", "28/02/2026", "Open"),
            ("Refund beneficiary revalidation", "Vertex Logistics LLP", "15/02/2026", "Open"),
        ],
    )

    doc.section("Financial impact")
    doc.table(
        ["Head", "Amount", "Notes"],
        [
            ("Service credits issued", "14,500", "12 customers"),
            ("Replacement goods", "62,300", "3 consignments"),
            ("Expedited freight", "18,900", "Vertex Logistics LLP"),
            ("Refunds reprocessed", "41,200", "2 customers"),
            ("Engineering effort", "96,000", "8 person-days"),
            ("Total", "2,32,900", ""),
        ],
    )

    doc.section("Customer identity verification")
    doc.para(
        "Identity was re-verified for escalated accounts. A specimen of the identity document submitted by "
        "the account holder for TKT-100302 is reproduced below.")
    doc.picture(_id_card("FATIMA BANO ANSARI", "Rehman Ansari", "21/07/1990", "AXWPA5521M",
                         "22 Rosewood Lane, Pune - 411 014"), 3.4)
    doc.para("Scan for the escalation dashboard:")
    doc.picture(_qr("https://sunrisetextiles.com/support/escalation/ESC-2026-0447"), 1.1)

    doc.section("Sign-off")
    doc.para(
        "Reviewed by Priya Nair on 2 February 2026. Internal contact: support@sunrisetextiles.com, "
        "+91 20 4505 3237.")
    doc.para("Authorised Signatory")
    doc.picture(_signature(4), 2.2)
    doc.para("Name: Priya Nair · Designation: Account Manager · For Sunrise Textiles Private Limited")
    doc.para(
        "Note on reference numbers: ticket, order and invoice numbers quoted throughout this file identify "
        "transactions rather than individuals and are retained deliberately. Policy revision 3.2 dated "
        "March 4, 2021 governs retention.")
    doc.save(path)


def main() -> int:
    build_ticket_log(HERE / "ticket_log.docx")
    build_offer_letter(HERE / "offer_letter.docx")
    build_claim_form(HERE / "claim_form.docx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
