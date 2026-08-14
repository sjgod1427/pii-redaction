"""Unit tests: run with `python -m pytest tests` (or `python tests/test_redactor.py`).

They cover the two things most likely to break silently — a detector losing a
PII shape, and the docx writer mangling a paragraph whose text is split across
several runs.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from pii_redactor import Policy
from pii_redactor.detectors import detect_dates_of_birth, detect_structured
from pii_redactor.addresses import detect_addresses
from pii_redactor.pipeline import Redactor
from pii_redactor.types import (
    ADDRESS, CREDIT_CARD, DOB, EMAIL, IP_ADDRESS, NATIONAL_ID, PHONE, SSN, URL,
)


def labels(spans) -> set[str]:
    return {s.label for s in spans}


def test_structured_types():
    text = (
        "Contact john.doe@example.com or +91 98765 43210. "
        "SSN 123-45-6789, card 4111 1111 1111 1111, host 10.0.12.7, PAN ABCDE1234F."
    )
    found = labels(detect_structured(text))
    assert {EMAIL, PHONE, SSN, CREDIT_CARD, IP_ADDRESS, NATIONAL_ID} <= found, found


def test_card_must_pass_luhn():
    spans = list(detect_structured("Order 4111 1111 1111 1112 shipped"))
    assert CREDIT_CARD not in labels(spans)


def test_ticket_and_order_numbers_are_not_pii():
    """Precision guard: bare reference numbers stay put."""
    text = "Ticket #4482910 relates to order 100002345 raised on invoice 993214."
    assert not list(detect_structured(text))


def test_dates_only_redacted_when_they_are_births():
    assert DOB not in labels(detect_dates_of_birth("The board met on March 4, 2021."))
    assert DOB in labels(detect_dates_of_birth("Date of birth: March 4, 1971"))


def test_address_span_covers_the_whole_address():
    text = "He lives at 12 Buena Monte, Panchvati, Pashan, Pune – 411 008, Maharashtra, India."
    spans = [s for s in detect_addresses(text) if s.label == ADDRESS]
    assert spans and spans[0].text.startswith("12 Buena Monte")
    assert spans[0].text.rstrip(".").endswith("India")


def test_address_does_not_swallow_surrounding_prose():
    text = (
        "Our Company, a company incorporated on July 30, 1979 under the Companies Act, "
        "having its Registered Office at 11/3, Village Birdewadi, Pune – 410 501, Maharashtra."
    )
    spans = [s for s in detect_addresses(text) if s.label == ADDRESS]
    assert spans and "incorporated" not in spans[0].text


def test_surrogates_are_deterministic_and_consistent():
    text = "Rashi Patil is our CFO. Rashi Patil can be reached at rashhi.patil@gmail.com."
    first = _redact(text)
    second = _redact(text)
    assert first == second, "same input + same seed must give the same output"
    assert "Rashi" not in first and "rashhi.patil" not in first
    # the same person is replaced by the same surrogate every time
    surname = first.split()[0]
    assert first.count(surname) >= 2


def test_docx_round_trip_preserves_runs(tmp_path=Path("/tmp")):
    import docx

    document = docx.Document()
    paragraph = document.add_paragraph()
    for fragment in ("Rashi ", "Pat", "il"):  # a name split across three runs
        paragraph.add_run(fragment)
    paragraph.add_run(" — +91 98765 43210")
    source = Path(tmp_path) / "_pii_test_input.docx"
    target = Path(tmp_path) / "_pii_test_output.docx"
    document.save(source)

    redactor = Redactor(Policy(disable_ner=True, extra_terms={"Rashi Patil": "PERSON"}))
    redactor.run(str(source), str(target))

    text = "\n".join(p.text for p in docx.Document(target).paragraphs)
    assert "Rashi" not in text and "98765" not in text
    assert len(docx.Document(target).paragraphs[0].runs) >= 3, "runs must survive"


# --- re-identification regressions -----------------------------------------
# Each of these is a defect found by auditing the redacted prospectus: every one
# left the issuer identifiable while the bulk-entity numbers looked healthy.


def test_url_split_across_runs_is_still_a_url():
    """Word breaks a long address over a line; the domain must still be caught."""
    spans = list(detect_structured("WEBSITE www.kshinternational. com"))
    assert URL in labels(spans), spans


def test_split_detection_does_not_invent_domains():
    """The gap-tolerant pattern must not join ordinary sentences into a host."""
    for text in (
        "Visit the site. Company said so.",
        "the Offer. Details on page 398.",
        "in Fiscal 2024. Revenue grew sharply.",
    ):
        found = labels(detect_structured(text))
        assert URL not in found and EMAIL not in found, (text, found)


def test_reference_codes_are_not_postal_codes():
    """"TKT-100294" is a ticket, not a PIN — the dash must separate words."""
    for text in ("3. Ticket TKT-100294 — payment declined",
                 "Policy NAV-HLT-4471209 renewed",
                 "Requisition REQ-2026-0091 approved"):
        assert not list(detect_addresses(text)), (text, list(detect_addresses(text)))
    assert list(detect_addresses("Office at 12 Main Road, Baner, Pune – 411 045")), "real PIN must still fire"


def test_address_does_not_swallow_a_sentence_containing_pii():
    """Prose around an address must stay out of the span — and keep its own labels."""
    text = ("The customer's SSN on file is 123-45-6789, recorded during the pilot, "
            "and his address is 940 Larch Street, Springfield, IL 62704, USA.")
    from pii_redactor.types import merge_spans

    merged = merge_spans(list(detect_structured(text)) + list(detect_addresses(text)))
    found = labels(merged)
    assert SSN in found, merged
    address = [s for s in merged if s.label == ADDRESS]
    assert address and "SSN" not in address[0].text, address


def test_bare_registration_number_is_pii():
    """The six-digit tail of a CIN identifies the company on its own."""
    spans = list(detect_structured("Registration number: 141032"))
    assert NATIONAL_ID in labels(spans), spans


def test_middle_initial_does_not_disqualify_a_name():
    from pii_redactor.gazetteer import Gazetteer

    gazetteer = Gazetteer(Policy(disable_ner=True))
    gazetteer.learn_table_row("Name | Holding", ["Rupal K. Sancheti", "39,062"])
    gazetteer.finalise()
    assert gazetteer.knows("Rupal K. Sancheti"), gazetteer.dropped


def test_place_names_are_learned_from_addresses_and_redacted_in_prose():
    """The design gap: a plant town in narrative prose reaches no detector."""
    from pii_redactor.gazetteer import Gazetteer
    from pii_redactor.types import LOCATION

    gazetteer = Gazetteer(Policy(disable_ner=True))
    gazetteer.learn_block(
        "Registered Office at Plot 4, Village Birdewadi, Chakan, Khed, Pune - 410 501, Maharashtra, India"
    )
    gazetteer.finalise()
    prose = "Our Chakan plant supplies the Pune facility."
    found = {s.text for s in gazetteer.place_spans(prose)}
    assert "Chakan" in found and "Pune" in found, (found, gazetteer.place_names)


def test_places_do_not_swallow_defined_terms_or_countries():
    from pii_redactor.gazetteer import Gazetteer

    gazetteer = Gazetteer(Policy(disable_ner=True))
    for junk in ("the Cap Price", "Designated RTA", "the United States", "N.A", "USD"):
        gazetteer._add_place(junk, source="test")
    assert not gazetteer.place_names, gazetteer.place_names


def test_postal_code_next_to_a_place_is_redacted():
    from pii_redactor.gazetteer import Gazetteer

    gazetteer = Gazetteer(Policy(disable_ner=True))
    gazetteer.learn_block("Office at 12 Main Road, Baner, Pune - 411 045, Maharashtra, India")
    gazetteer.finalise()
    spans = list(gazetteer.place_spans("Pune 411 045 Maharashtra, India"))
    assert any("411 045" in s.text for s in spans), spans


def test_surrogate_addresses_do_not_reuse_real_state_names():
    """A fake address ending in a real state puts back what LOCATION removed."""
    from pii_redactor.surrogates import INDIAN_STATE_POOL
    from pii_redactor.config import INDIAN_STATES

    for name in INDIAN_STATE_POOL:
        assert name.casefold() not in INDIAN_STATES, name


def test_image_audit_summary_never_quotes_recovered_text():
    """detections.csv must not undo the image redaction it is auditing."""
    from pii_redactor.images import ImageEvidence

    # Invented values: a test for not disclosing identifiers must not itself
    # embed real ones lifted from the document under test.
    evidence = ImageEvidence(
        name="card.png", width=600, height=400, format="PNG",
        ocr_text="ARJUN TESTWALLA\nFather: Ramesh Testwalla\nDOB 03/09/1977",
        codes=["https://example.org/secret"],
    )
    summary = evidence.summary()
    for secret in ("ARJUN", "Testwalla", "03/09/1977", "secret"):
        assert secret not in summary, summary


# --- images ----------------------------------------------------------------
# Fixtures are generated here rather than read from the prospectus, so the image
# stage is tested on documents the tool has never seen.


def _image_redactor(**policy_kwargs):
    from pii_redactor.images import ImageRedactor

    redactor = Redactor(Policy(disable_ner=True, **policy_kwargs))
    redactor.gazetteer.finalise()
    redactor.surrogates.bind(redactor.gazetteer)
    return ImageRedactor(redactor.policy, redactor.detect_text, redactor.surrogates.replacement)


def _png(draw_on, size=(420, 260), colour="white") -> bytes:
    import io

    from PIL import Image, ImageDraw

    image = Image.new("RGB", size, colour)
    draw_on(ImageDraw.Draw(image), size)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def test_image_carrying_an_identifier_is_called_an_identity_document():
    from pii_redactor.types import IMAGE_ID_DOCUMENT

    def card(draw, size):
        for index, line in enumerate(
            ["INCOME TAX DEPARTMENT", "Permanent Account Number", "ABCDE1234F",
             "Name: JOHN QUINCY", "Date of Birth 06/05/1988", "Signature"]
        ):
            draw.text((14, 20 + index * 34), line, fill="black")

    redactor = _image_redactor()
    evidence = redactor.analyse(_png(card), "card.png")
    label, reason = redactor.classify(evidence, context="")
    assert label == IMAGE_ID_DOCUMENT, (label, reason, evidence.ocr_text)


def test_signature_is_caught_by_its_caption():
    from pii_redactor.types import IMAGE_SIGNATURE

    redactor = _image_redactor()
    evidence = redactor.analyse(_png(lambda draw, size: draw.line([(20, 150), (380, 120)], fill="navy", width=3)), "sig.png")
    label, _ = redactor.classify(evidence, context="For and on behalf of the Board\nSignature")
    assert label == IMAGE_SIGNATURE, label


def test_images_are_redacted_when_no_analysis_engine_is_installed():
    """A slim deployment without OCR must still replace pictures, not keep them.

    Without OpenCV the ink measurement silently reports 0.0 — not because the
    picture is blank but because nothing looked at it — which used to route
    every image to the "blank, keep it" branch.
    """
    from pii_redactor.types import IMAGE_UNCLASSIFIED

    redactor = _image_redactor()
    redactor.engines._loaded = {"ocr": True, "cv": True}
    redactor.engines._ocr = None
    redactor.engines.cv2 = None
    assert not redactor.engines.available

    evidence = redactor.analyse(_png(lambda draw, size: None), "x.png")
    label, reason = redactor.classify(evidence, context="")
    assert label == IMAGE_UNCLASSIFIED, (label, reason)


def test_unreadable_image_is_redacted_not_trusted():
    """Fail safe: something that cannot be inspected cannot be cleared."""
    from pii_redactor.types import IMAGE_UNCLASSIFIED

    redactor = _image_redactor()
    evidence = redactor.analyse(b"\x89PNG\r\n\x1a\n corrupted payload", "broken.png")
    assert not evidence.readable
    label, _ = redactor.classify(evidence, context="")
    assert label == IMAGE_UNCLASSIFIED, label


def test_blank_image_is_left_alone_but_unclassified_ink_is_not():
    redactor = _image_redactor()
    blank = redactor.analyse(_png(lambda draw, size: None), "blank.png")
    assert redactor.classify(blank, context="")[0] == "", "a blank image should be kept"

    def scribble(draw, size):
        draw.rectangle([(10, 10), (300, 200)], fill="#3a3a3a")

    inked = redactor.analyse(_png(scribble), "ink.png")
    assert redactor.classify(inked, context="")[0] != "", "ink with no evidence must be redacted"


def test_keep_unclassified_images_flips_the_default():
    redactor = _image_redactor(keep_unclassified_images=True)

    def scribble(draw, size):
        draw.rectangle([(10, 10), (300, 200)], fill="#3a3a3a")

    evidence = redactor.analyse(_png(scribble), "ink.png")
    assert redactor.classify(evidence, context="")[0] == ""


def test_tiny_icons_are_not_touched():
    redactor = _image_redactor()
    evidence = redactor.analyse(_png(lambda draw, size: None, size=(12, 12), colour="black"), "bullet.png")
    assert redactor.classify(evidence, context="")[0] == ""


def test_replacement_keeps_dimensions_and_format():
    from pii_redactor.types import IMAGE_ID_DOCUMENT

    redactor = _image_redactor()
    evidence = redactor.analyse(_png(lambda draw, size: None, size=(321, 197)), "x.png")
    encoded = redactor.encode(redactor.render(IMAGE_ID_DOCUMENT, evidence, ""), evidence)
    import io

    from PIL import Image

    with Image.open(io.BytesIO(encoded)) as replacement:
        assert replacement.size == (321, 197)
        assert replacement.format == "PNG"


def test_url_path_does_not_keep_the_customer_name():
    """A surrogate host is not enough if the path still names the person."""
    out = _redact("Profile at https://acme.com/account/rashi-patil?ref=Patil for review.")
    assert "rashi" not in out.lower() and "patil" not in out.lower(), out


def test_redacted_qr_encodes_a_surrogate_url():
    """The replacement still scans — to a fake address, consistently with the text."""
    import io

    import cv2
    import numpy
    import qrcode
    from PIL import Image

    from pii_redactor.types import IMAGE_CODE

    buffer = io.BytesIO()
    qrcode.make("https://realcompany.example.org/secret").save(buffer, format="PNG")
    redactor = _image_redactor()
    evidence = redactor.analyse(buffer.getvalue(), "qr.png")
    assert evidence.codes, "the fixture QR should decode"

    encoded = redactor.encode(redactor.render(IMAGE_CODE, evidence, ""), evidence)
    with Image.open(io.BytesIO(encoded)) as replacement:
        array = numpy.array(replacement.convert("RGB"))[:, :, ::-1].copy()
    decoded = cv2.QRCodeDetector().detectAndDecode(array)[0]
    assert decoded and "realcompany" not in decoded, decoded


def _redact(text: str) -> str:
    redactor = Redactor(Policy(disable_ner=True, extra_terms={"Rashi Patil": "PERSON"}))
    redactor.gazetteer.add_manual(redactor.policy.extra_terms)
    redactor.gazetteer.finalise()
    redactor.surrogates.bind(redactor.gazetteer)
    spans = redactor.detect(text)
    out, cursor = [], 0
    for span in spans:
        out.append(text[cursor : span.start])
        out.append(redactor.surrogates.replacement(span))
        cursor = span.end
    out.append(text[cursor:])
    return "".join(out)


if __name__ == "__main__":
    failures = 0
    for name, function in sorted(globals().items()):
        if name.startswith("test_") and callable(function):
            try:
                function()
                print(f"PASS {name}")
            except AssertionError as error:
                failures += 1
                print(f"FAIL {name}: {error}")
    raise SystemExit(1 if failures else 0)
